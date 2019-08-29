from django.http import HttpResponse
from django.utils.datetime_safe import datetime
from docx import Document
from django.shortcuts import get_object_or_404, render
from django.views.decorators.http import require_GET, require_POST
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm

from chair.forms import FilterReviewsForm
from chair.utility import validate_chair_access, build_paged_view_context
from conferences.models import Conference
from review.forms import UpdateDecisionForm
from review.models import Decision, ReviewStats
from review.utilities import get_average_score, qualify_score, \
    count_required_reviews, UNKNOWN_QUALITY, count_missing_reviews
from submissions.models import Submission
from users.models import User


def _render_feed_item(request, submission, conference):
    stats, _ = ReviewStats.objects.get_or_create(conference=conference)

    # Fill reviews data - a list of scores with data, and warnings:
    reviews_data = []
    num_incomplete, num_missing = 0, 0
    for review in submission.reviews.all():
        score = review.average_score()
        quality = qualify_score(score, stats=stats)
        reviews_data.append({
            'quality': quality,
            'score': f'{score:.1f}' if score > 0 else '?',
        })
        if score == 0:
            num_incomplete += 1
    num_required = count_required_reviews(submission)
    if num_required > len(reviews_data):
        num_missing = num_required - len(reviews_data)
        for _ in range(num_required - len(reviews_data)):
            reviews_data.append({'quality': UNKNOWN_QUALITY, 'score': '-'})

    warnings = []
    if num_incomplete > 0:
        warnings.append(f'{num_incomplete} reviews are not finished')
    if num_missing > 0:
        warnings.append(f'{num_missing} reviews are not assigned')

    return render(request, 'chair/reviews/_submission_feed_item.html', context={
        'submission': submission,
        'conf_pk': conference.pk,
        'decision_data': _get_decision_form_data(submission),
        'reviews_data': reviews_data,
        'warnings': warnings,
    })


def _get_decision_form_data(submission):
    decision = submission.review_decision.first()
    proc_type = decision.proc_type if decision else None
    volume = decision.volume if decision else None
    default_option = [('', 'Not selected')]

    # 1) Filling data_decision value and display:
    data_decision = {'hidden': False, 'options': Decision.DECISION_CHOICES}
    decision_value = Decision.UNDEFINED if not decision else decision.decision
    data_decision['value'] = decision_value
    data_decision['display'] = [
        opt[1] for opt in Decision.DECISION_CHOICES if opt[0] == decision_value
    ][0]

    # 2) Fill proceedings type if possible:
    stype = submission.stype
    data_proc_type = {
        'hidden': decision_value != Decision.ACCEPT,
        'value': '', 'display': default_option[0][-1],
    }
    if not data_proc_type['hidden']:
        data_proc_type['options'] = default_option + [
            (pt.pk, pt.name) for pt in stype.possible_proceedings.all()]
        if proc_type:
            data_proc_type.update({
                'value': proc_type.pk, 'display': proc_type.name
            })

    # 3) Fill volumes if possible:
    data_volume = {
        'hidden': data_proc_type['value'] == '',
        'value': '', 'display': default_option[0][-1],
    }
    if not data_volume['hidden']:
        data_volume['options'] = default_option + [
            (vol.pk, vol.name) for vol in proc_type.volumes.all()]
        if volume:
            data_volume.update({'value': volume.pk, 'display': volume.name})

    # 4) Collect everything and output:
    return {
        'decision': data_decision,
        'proc_type': data_proc_type,
        'volume': data_volume,
        'committed': decision.committed if decision else True
    }


@require_GET
def list_submissions(request, conf_pk, page=1):
    conference = get_object_or_404(Conference, pk=conf_pk)
    validate_chair_access(request.user, conference)
    submissions = conference.submission_set.exclude(status=Submission.SUBMITTED)

    filter_form = FilterReviewsForm(request.GET, instance=conference)
    if filter_form.is_valid():
        submissions = filter_form.apply(submissions)

    cached_stypes = {st.pk: st for st in conference.submissiontype_set.all()}
    scores = [{
        'pk': sub.pk, 'score': get_average_score(sub),
        'num_missing': count_missing_reviews(sub, cached_stypes=cached_stypes)
    } for sub in submissions]

    if scores:
        max_pk = max([record['pk'] for record in scores])

        def sort_key(score_record):
            if score_record['score'] > 0 and score_record['num_missing'] == 0:
                return score_record['score']
            pk_diff = score_record['pk'] / max_pk  # always in 0..1!
            return -score_record['num_missing'] - pk_diff

        scores.sort(key=sort_key, reverse=True)

    pks = [score_record['pk'] for score_record in scores]
    context = build_paged_view_context(
        request, pks, page, 'chair:reviews-pages', {'conf_pk': conf_pk})
    context.update({'conference': conference, 'filter_form': filter_form})
    return render(request, 'chair/reviews/reviews_list.html', context=context)


def decision_control_panel(request, conf_pk, sub_pk):
    conference = get_object_or_404(Conference, pk=conf_pk)
    validate_chair_access(request.user, conference)
    submission = Submission.objects.get(pk=sub_pk)
    decision = submission.review_decision.first()
    if not decision:
        decision = Decision.objects.create(submission=submission)
    form = UpdateDecisionForm(request.POST or None, instance=decision)
    if request.method == 'POST' and form.is_valid():
        form.save()
    return _render_feed_item(request, submission, conference)


@require_POST
def commit_decision(request, conf_pk, sub_pk):
    conference = get_object_or_404(Conference, pk=conf_pk)
    validate_chair_access(request.user, conference)
    submission = Submission.objects.get(pk=sub_pk)
    decision = submission.review_decision.first()
    decision.commit()
    return _render_feed_item(request, submission, conference)


@require_GET
def list_item(request, conf_pk, sub_pk):
    conference = get_object_or_404(Conference, pk=conf_pk)
    validate_chair_access(request.user, conference)
    submission = Submission.objects.get(pk=sub_pk)
    return _render_feed_item(request, submission, conference)


# @require_GET
# def export_doc(request, conf_pk):
#     conference = get_object_or_404(Conference, pk=conf_pk)
#     validate_chair_access(request.user, conference)
#     submissions = conference.submission_set.exclude(status=Submission.SUBMITTED)
#     stats, _ = ReviewStats.objects.get_or_create(conference=conference)
#     timestamp = datetime.now().strftime("%Y%m%d_%H%M")
#
#     document = Document()
#     document.add_heading(f'{conference.short_name} Review Report')
#     document.add_paragraph('')  # Just line skip
#     table = document.add_table(rows=9, cols=2)
#     table.rows[0].cells[0].text = 'Report date'
#     table.rows[0].cells[1].text = datetime.now().strftime('%d %b %Y, %H:%M')
#     for i, (key, details, dtype) in enumerate(zip(
#             ('average', 'lq', 'median', 'hq',
#              'num_review_submissions', 'num_complete',
#              'num_assigned_incomplete', 'num_not_assigned'),
#             ('Average score', 'Q3', 'Q2 (median)', 'Q1',
#              'Number of submissions under review',
#              'Number of submissions with complete reviews',
#              'Number of submissions with assigned, but incomplete reviews',
#              'Number of submissions with partially not assigned reviewers'),
#             ('float', 'float', 'float', 'float', 'int', 'int', 'int', 'int')
#     )):
#         table.rows[i + 1].cells[0].text = details
#         value = f'{stats[key]:.2f}' if dtype == 'float' else f'{stats[key]}'
#         table.rows[i + 1].cells[1].text = value
#
#     # Writing submissions:
#     submissions = list(submissions)
#
#     def get_order_key(sub):
#         rs = reviews[sub]
#         if rs['complete']:
#             return rs['score']
#         return -(rs['num_required'] - rs['num_finished'])
#
#     submissions.sort(key=get_order_key, reverse=True)
#
#     profiles = {u: u.profile for u in User.objects.all()}
#     document.add_page_break()
#
#     for submission in submissions:
#         record = reviews[submission]
#         scores_str_list = [
#             f'{sc:.1f}' if sc != '?' else '?' for sc in record['scores_list']]
#         try:
#             document.add_heading(
#                 f'#{submission.pk}: {submission.title}', level=1)
#         except ValueError:
#             document.add_heading(
#                 f'#{submission.pk}: [title hidden due to illegal characters]',
#                 level=1)
#
#         p = document.add_paragraph()
#         p.add_run('Review Score: ').bold = True
#         p.add_run(f'{record["score"]:.2f} ({record["quality"]}, scores: '
#                   f'{"/".join(scores_str_list)})')
#         p = document.add_paragraph()
#         p.add_run('Reviews finished / assigned / required: ').bold = True
#         p.add_run(f"{record['num_finished']} / {record['num_assigned']} / "
#                   f"{record['num_required']}")
#         p = document.add_paragraph()
#         p.add_run('Authors: ').bold = True
#         authors = submission.authors.all()
#         for i, author in enumerate(authors):
#             profile = profiles[author.user]
#             p = document.add_paragraph(f'{i + 1}. ')
#             p.add_run(profile.get_full_name()).bold = True
#             name_rus = f'{profile.first_name_rus} {profile.last_name_rus}'
#             if name_rus != ' ':
#                 p.add_run(f' [{name_rus}]')
#             p.add_run(' (')
#             p.add_run(
#                 f'{profile.get_country_display()}, {profile.affiliation}, '
#                 f'{profile.degree}').italic = True
#             p.add_run(')')
#         # p = document.add_paragraph()
#         # p.add_run('Abstract:').bold = True
#         # document.add_paragraph(submission.abstract)
#         for i, review in enumerate(submission.reviews.all()):
#             reviewer_profile = profiles[review.reviewer.user]
#             document.add_heading(
#                 f'Review #{i+1} by {reviewer_profile.get_full_name()}',
#                 level=2)
#             review_data = (
#                 ('Technical merit', review.technical_merit),
#                 ('Originality', review.originality),
#                 ('Relevance', review.relevance),
#                 ('Clarity', review.clarity),
#                 ('Finished', 'Yes' if review.submitted else 'No'),
#             )
#             table = document.add_table(rows=len(review_data), cols=2)
#             for row_i, row in enumerate(review_data):
#                 table.rows[row_i].cells[0].text = row[0]
#                 table.rows[row_i].cells[1].text = row[1]
#             for row in table.rows:
#                 row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
#                 row.height = Cm(0.7)
#             try:
#                 document.add_paragraph(review.details)
#             except ValueError:
#                 p = document.add_paragraph()
#                 p.add_run('[Review details hidden since they contain illegal '
#                           'characters and can not be processed in .docx '
#                           'format]').italic = True
#
#     response = HttpResponse(
#         content_type='application/vnd.openxmlformats-officedocument.'
#                      'wordprocessingml.document')
#     response['Content-Disposition'] = \
#         f'attachment; filename=reviews-{timestamp}.docx'
#     document.save(response)
#
#     return response
#
