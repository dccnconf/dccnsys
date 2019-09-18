import csv
from datetime import datetime

from django.db.models import F, Count, Q, Value, CharField
from django.db.models.functions import Concat
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from django.urls import reverse
from django.views.decorators.http import require_GET
from django_countries import countries
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm

from chair.forms import ExportSubmissionsForm, ExportUsersForm
from conferences.utilities import validate_chair_access
from conferences.models import Conference
from review.models import ReviewStats, Review
from review.utilities import get_average_score
from submissions.models import Submission
from users.models import User


def create_export_view(form_class, file_name_prefix, title):

    def view(request, conf_pk):
        conference = get_object_or_404(Conference, pk=conf_pk)
        validate_chair_access(request.user, conference)

        default_next = reverse('chair:home', kwargs={'conf_pk': conf_pk})
        next_url = request.GET.get('next', default_next)

        if request.method == 'POST':
            form = form_class(request.POST, conference=conference)
            if form.is_valid():
                data = form.apply(request)

                # Create the HttpResponse object with the appropriate header.
                response = HttpResponse(content_type='text/csv')
                timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                response['Content-Disposition'] = \
                    f'attachment; filename="{file_name_prefix}-{timestamp}.csv"'
                writer = csv.writer(response)

                # import sys
                # writer = csv.writer(sys.stdout)
                writer.writerow(form.cleaned_data['columns'])
                for item in data:
                    row = []
                    for col in form.cleaned_data['columns']:
                        row.append(item[col])
                    writer.writerow(row)
                return response
        else:  # request was GET:
            form = form_class(conference=conference)

        return render(request, 'chair/export/export_dialog.html', context={
            'conference': conference,
            'next': next_url,
            'form': form,
            'panel_title': title,
        })

    return view


export_submissions = create_export_view(
    ExportSubmissionsForm, 'papers', 'Export submissions')

export_users = create_export_view(ExportUsersForm, 'users', 'Export users')


@require_GET
def export_reviews_doc(request, conf_pk):
    conference = get_object_or_404(Conference, pk=conf_pk)
    validate_chair_access(request.user, conference)

    stats, _ = ReviewStats.objects.get_or_create(conference=conference)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")

    #
    # Create document, write title and key statistics:
    #
    document = Document()
    document.add_heading(f'{conference.short_name} Review Report')
    document.add_paragraph('')
    table = document.add_table(rows=9, cols=2)
    table.rows[0].cells[0].text = 'Report date'
    table.rows[0].cells[1].text = datetime.now().strftime('%d %b %Y, %H:%M')
    for i, (key, details, dtype) in enumerate(zip(
            ('average_score', 'q1_score', 'median_score', 'q3_score',
             'num_submissions_reviewed',
             'num_submissions_with_incomplete_reviews',
             'num_submissions_with_missing_reviewers'),
            ('Average score', 'Q1 (lowest)', 'Q2 (median)', 'Q3 (highest)',
             'Number of submissions reviewed',
             'Number of submissions with incomplete reviews',
             'Number of submissions missing one or more reviewers'),
            ('float', 'float', 'float', 'float', 'int', 'int', 'int')
    )):
        table.rows[i + 1].cells[0].text = details
        if dtype == 'float':
            value = f'{getattr(stats, key, 0.0):.2f}'
        elif dtype == 'int':
            value = f'{getattr(stats, key, 0):d}'
        else:
            value = getattr(stats, key, '')
        table.rows[i + 1].cells[1].text = value
    document.add_page_break()

    #
    # Write submissions and their reviews:
    #
    submissions = conference.submission_set.filter(status__in={
        Submission.UNDER_REVIEW, Submission.ACCEPTED, Submission.REJECTED,
        Submission.IN_PRINT, Submission.PUBLISHED}
    ).annotate(
        num_reviews_required=F('stype__num_reviews'),
        num_reviews_assigned=Count('reviews', distinct=True),
        num_reviews_submitted=Count('reviews', filter=Q(
            reviews__submitted=True), distinct=True)
    ).order_by('pk').distinct()

    users = User.objects.annotate(
        full_name=Concat(
            'profile__last_name', Value(' '), 'profile__first_name',
            output_field=CharField()),
        full_name_rus=Concat(
            'profile__last_name_rus', Value(' '), 'profile__first_name_rus',
            Value(' '), 'profile__middle_name_rus', output_field=CharField()),
        country=F('profile__country'),
        affiliation=F('profile__affiliation'),
        degree=F('profile__degree'),
    ).in_bulk()

    scores = {sub.pk: get_average_score(sub) for sub in submissions}
    countries_dict = dict(countries)

    for submission in submissions:
        try:
            document.add_heading(
                f'#{submission.pk}: {submission.title}', level=1)
        except ValueError:
            document.add_heading(
                f'#{submission.pk}: [title hidden due to illegal characters',
                level=1)

        p = document.add_paragraph()
        p.add_run('Status: ').bold = True
        p.add_run(f'{submission.get_status_display()}')

        p = document.add_paragraph()
        p.add_run('Review Score: ').bold = True
        p.add_run(f'{scores[submission.pk]:.2f}')

        p = document.add_paragraph()
        p.add_run('Reviews finished / assigned / required: ').bold = True
        p.add_run(f'{submission.num_reviews_submitted} / '
                  f'{submission.num_reviews_assigned} / '
                  f'{submission.num_reviews_required}')

        p = document.add_paragraph()
        p.add_run('Authors: ').bold = True
        for i, author in enumerate(submission.authors.all()):
            user = users[author.user_id]

            p = document.add_paragraph(f'{i+1}. ')
            p.add_run(user.full_name)
            rus_name = user.full_name_rus.strip()
            if rus_name:
                p.add_run(f' [{rus_name}]')

            country = countries_dict[user.country]
            p.add_run(
                f' ({country}, {user.affiliation}, {user.degree})'
            ).italic = True

        p = document.add_paragraph()
        p.add_run('Abstract: ').bold = True
        document.add_paragraph(submission.abstract)

        for i, review in enumerate(Review.objects.filter(paper=submission)):
            user = users[review.reviewer.user_id]
            document.add_heading(f'Review #{i+1} by {user.full_name}', level=2)
            review_data = (
                ('Technical merit', review.technical_merit),
                ('Originality', review.originality),
                ('Relevance', review.relevance),
                ('Clarity', review.clarity),
                ('Finished', 'Yes' if review.submitted else 'No')
            )
            table = document.add_table(rows=len(review_data), cols=2)
            for row_i, row in enumerate(review_data):
                table.rows[row_i].cells[0].text = row[0]
                table.rows[row_i].cells[1].text = row[1]
            for row in table.rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = Cm(0.7)

            try:
                document.add_paragraph(review.details)
            except ValueError:
                p = document.add_paragraph()
                p.add_run(
                    '[Review details hidden since they contain illegal '
                    'characters and can not be processed in DOC-export]'
                ).italic = True

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.'
                     'wordprocessingml.document')
    response['Content-Disposition'] = \
        f'attachment; filename=reviews-{timestamp}.docx'
    document.save(response)

    return response
